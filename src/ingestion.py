"""
ingestion.py — Document ingestion pipeline.

Supports:
  • PDF files (pypdf / pymupdf)
  • DOCX files (python-docx)
  • Plain text files (.txt, .md)
  • Web URLs (trafilatura for clean text extraction)

Returns a list of LangChain Document objects with rich metadata.
"""
from __future__ import annotations
import hashlib
import tempfile
from pathlib import Path
from typing import List, Optional
from datetime import datetime

import requests
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

from src.config import settings
from src.logger import logger


# ── Text splitter ──────────────────────────────────────────────────────────────
def _get_splitter() -> RecursiveCharacterTextSplitter:
    return RecursiveCharacterTextSplitter(
        chunk_size=settings.chunk_size,
        chunk_overlap=settings.chunk_overlap,
        separators=["\n\n", "\n", ". ", "! ", "? ", " ", ""],
        length_function=len,
    )


def _make_doc_id(content: str) -> str:
    return hashlib.md5(content.encode()).hexdigest()[:12]


# ── PDF Loader ─────────────────────────────────────────────────────────────────
def load_pdf(file_path: str | Path, source_name: Optional[str] = None) -> List[Document]:
    """Load a PDF and split into chunks."""
    file_path = Path(file_path)
    source_name = source_name or file_path.name
    logger.info(f"Loading PDF: {file_path.name}")

    try:
        import fitz  # pymupdf — better quality extraction
        pdf = fitz.open(str(file_path))
        pages = []
        for i, page in enumerate(pdf):
            text = page.get_text("text")
            if text.strip():
                pages.append(Document(
                    page_content=text,
                    metadata={
                        "source": source_name,
                        "source_type": "pdf",
                        "page": i + 1,
                        "total_pages": len(pdf),
                        "ingested_at": datetime.now().isoformat(),
                    }
                ))
        pdf.close()
    except Exception:
        # Fallback to pypdf
        logger.warning("PyMuPDF failed, falling back to pypdf")
        from pypdf import PdfReader
        reader = PdfReader(str(file_path))
        pages = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            if text.strip():
                pages.append(Document(
                    page_content=text,
                    metadata={
                        "source": source_name,
                        "source_type": "pdf",
                        "page": i + 1,
                        "total_pages": len(reader.pages),
                        "ingested_at": datetime.now().isoformat(),
                    }
                ))

    splitter = _get_splitter()
    chunks = splitter.split_documents(pages)
    logger.info(f"PDF '{source_name}' → {len(pages)} pages → {len(chunks)} chunks")
    return chunks


# ── DOCX Loader ────────────────────────────────────────────────────────────────
def load_docx(file_path: str | Path, source_name: Optional[str] = None) -> List[Document]:
    """Load a .docx file and split into chunks."""
    from docx import Document as DocxDocument
    file_path = Path(file_path)
    source_name = source_name or file_path.name
    logger.info(f"Loading DOCX: {file_path.name}")

    doc = DocxDocument(str(file_path))
    full_text = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())

    base_doc = Document(
        page_content=full_text,
        metadata={
            "source": source_name,
            "source_type": "docx",
            "ingested_at": datetime.now().isoformat(),
        }
    )
    splitter = _get_splitter()
    chunks = splitter.split_documents([base_doc])
    logger.info(f"DOCX '{source_name}' → {len(chunks)} chunks")
    return chunks


# ── TXT / MD Loader ────────────────────────────────────────────────────────────
def load_text_file(file_path: str | Path, source_name: Optional[str] = None) -> List[Document]:
    """Load a plain-text or markdown file."""
    file_path = Path(file_path)
    source_name = source_name or file_path.name
    logger.info(f"Loading text file: {file_path.name}")

    text = file_path.read_text(encoding="utf-8", errors="replace")
    base_doc = Document(
        page_content=text,
        metadata={
            "source": source_name,
            "source_type": "text",
            "ingested_at": datetime.now().isoformat(),
        }
    )
    splitter = _get_splitter()
    chunks = splitter.split_documents([base_doc])
    logger.info(f"Text file '{source_name}' → {len(chunks)} chunks")
    return chunks


# ── Web URL Loader ─────────────────────────────────────────────────────────────
def load_url(url: str) -> List[Document]:
    """Scrape a URL using trafilatura for clean article extraction."""
    import trafilatura
    logger.info(f"Fetching URL: {url}")

    try:
        headers = {"User-Agent": "Mozilla/5.0 (research-assistant-bot/1.0)"}
        resp = requests.get(url, timeout=15, headers=headers)
        resp.raise_for_status()
        html = resp.text
    except Exception as e:
        logger.error(f"Failed to fetch {url}: {e}")
        raise RuntimeError(f"Could not fetch URL: {url}\n{e}")

    text = trafilatura.extract(
        html,
        include_comments=False,
        include_tables=True,
        no_fallback=False,
    )
    if not text or len(text.strip()) < 100:
        # Fallback: BeautifulSoup
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(html, "html.parser")
        for tag in soup(["script", "style", "nav", "footer", "header"]):
            tag.decompose()
        text = soup.get_text(separator="\n", strip=True)

    if not text:
        raise RuntimeError(f"No readable content extracted from {url}")

    # Extract title
    from bs4 import BeautifulSoup
    soup = BeautifulSoup(html, "html.parser")
    title = soup.title.string.strip() if soup.title else url

    base_doc = Document(
        page_content=text,
        metadata={
            "source": url,
            "source_type": "web",
            "title": title,
            "ingested_at": datetime.now().isoformat(),
        }
    )
    splitter = _get_splitter()
    chunks = splitter.split_documents([base_doc])
    logger.info(f"URL '{url}' → {len(chunks)} chunks")
    return chunks


# ── Dispatcher ─────────────────────────────────────────────────────────────────
def ingest_file(file_bytes: bytes, filename: str) -> List[Document]:
    """
    Universal entry point for Streamlit file uploads.
    Saves bytes to a temp file, then dispatches to the right loader.
    """
    suffix = Path(filename).suffix.lower()
    with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
        tmp.write(file_bytes)
        tmp_path = Path(tmp.name)

    try:
        if suffix == ".pdf":
            return load_pdf(tmp_path, source_name=filename)
        elif suffix in (".docx", ".doc"):
            return load_docx(tmp_path, source_name=filename)
        elif suffix in (".txt", ".md", ".rst"):
            return load_text_file(tmp_path, source_name=filename)
        else:
            raise ValueError(f"Unsupported file type: {suffix}")
    finally:
        tmp_path.unlink(missing_ok=True)
